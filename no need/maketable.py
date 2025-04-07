import pandas as pd
import numpy as np
import argparse
from docx import Document

def round_if_float(x):
    try:
        return round(float(x), 2)
    except (ValueError, TypeError):
        return x
    
def round_ifNe_float(x):
    try:
        return round(float(x))
    except (ValueError, TypeError):
        return x
    
def change_if_only(x):
    if x=='±':
        return '-'
    else:
        return x
    
def makeDF(data,ITYPE,VTYPE,VSE,rownametemp,colnametemp):
    data.loc[(data['denSNP']==0),'imp']="ped"
    data.loc[(data['denSNP']==1250),'imp']="hd"
    if ITYPE=='T':
        data.loc[data['denSNP']==0,VTYPE] = ''
        data.loc[data['denSNP']==0,VSE] = ''
        data.loc[data['denSNP']==1250,VTYPE] = ''
        data.loc[data['denSNP']==1250,VSE] = ''
    data_filt = data[(data['imp'].isin([ITYPE,"ped","hd"]) ) & (data['nGeneration']==20)]
    z_matrix_value = data_filt.pivot(index='denSNP', columns='nRef', values=VTYPE)
    z_matrix_se = data_filt.pivot(index='denSNP', columns='nRef', values=VSE)
    if VTYPE in ('genicNe','Ne'):
        z_matrix_value = z_matrix_value.applymap(round_ifNe_float)
        z_matrix_se = z_matrix_se.applymap(round_ifNe_float)
    else:
        z_matrix_value = z_matrix_value.applymap(round_if_float)
        z_matrix_se = z_matrix_se.applymap(round_if_float)
    df1_str = z_matrix_value.applymap(str)
    df2_str = z_matrix_se.applymap(str)
    df_combined = df1_str + '±' + df2_str
    df_combined.index = rownametemp
    df_combined.columns = colnametemp
    df_combined = df_combined.applymap(change_if_only)
    return df_combined

def process_data(VTYPE,VSE,TNAME):
    data = pd.read_csv("allzt.csv")
    colnames_temp = sorted(data['nRef'].unique())
    rownames_temp = sorted(data['denSNP'].unique())
    colnames_temp = [str(element) for element in colnames_temp]
    rownames_temp = [str(element) for element in rownames_temp]
    DF_F = makeDF(data=data,ITYPE="F",VTYPE=VTYPE,VSE=VSE,rownametemp=rownames_temp,colnametemp=colnames_temp)
    DF_T = makeDF(data=data,ITYPE="T",VTYPE=VTYPE,VSE=VSE,rownametemp=rownames_temp,colnametemp=colnames_temp)
    DF_dombined = DF_F + ' / ' + DF_T
    
    # 创建一个新的 Word 文档
    doc = Document()

    # 添加标题
    #doc.add_heading('Z_matrix Table', level=1)

    # 添加表格
    table = doc.add_table(rows=1, cols=len(DF_dombined.columns)+1)

    # 添加列标题
    hdr_cells = table.rows[0].cells
    for i, col_name in enumerate(DF_dombined.columns):
        hdr_cells[i+1].text = col_name

    # 添加行数据
    for index, row in DF_dombined.iterrows():
        row_cells = table.add_row().cells
        row_cells[0].text = index
        for i, cell_value in enumerate(row):
            row_cells[i+1].text = cell_value

    # 保存文档
    doc.save(TNAME)

if __name__ == "__main__":
    parser = argparse.ArgumentParser(description='Process some strings.')
    parser.add_argument('--VTYPE', type=str, help='Colnames of the value')
    parser.add_argument('--VSE', type=str, help='Colnames of the value error')
    parser.add_argument('--TNAME', type=str, help='The name of the out Table')

    args = parser.parse_args()

    # 调用函数并输出结果到文件
    process_data(args.VTYPE, args.VSE, args.TNAME)
